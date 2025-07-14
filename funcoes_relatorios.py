#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
📊 MÓDULO DE GERAÇÃO DE RELATÓRIOS
==================================

Funcionalidades para gerar relatórios pedagógicos e financeiros
usando OpenAI para formatação inteligente e python-docx para .docx
"""

import os
import io
import base64
from datetime import datetime, date
from typing import Dict, List, Optional, Union
import pandas as pd

# Dependências necessárias
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# Importar funções do modelo pedagógico
from models.pedagogico import (
    listar_turmas_disponiveis,
    obter_mapeamento_turmas,
    buscar_alunos_por_turmas,
    supabase
)

# ==========================================================
# 🔧 CONFIGURAÇÕES E CONSTANTES
# ==========================================================

# Configuração OpenAI
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
openai_client = None

if OPENAI_AVAILABLE and OPENAI_API_KEY:
    try:
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception as e:
        print(f"❌ Erro ao inicializar OpenAI: {e}")

# Campos disponíveis para relatórios
CAMPOS_ALUNO = {
    'nome': 'Nome do Aluno',
    'turno': 'Turno',
    'data_nascimento': 'Data de Nascimento',
    'dia_vencimento': 'Dia de Vencimento',
    'data_matricula': 'Data de Matrícula',
    'valor_mensalidade': 'Valor da Mensalidade',
    'situacao': 'Situação',
    'mensalidades_geradas': 'Mensalidades geradas?'
}

CAMPOS_RESPONSAVEL = {
    'nome': 'Nome do Responsável',
    'cpf': 'CPF',
    'telefone': 'Telefone/Contato',
    'email': 'Email',
    'endereco': 'Endereço',
    'tipo_relacao': 'Tipo de Relação',
    'responsavel_financeiro': 'Responsável Financeiro'
}

CAMPOS_MENSALIDADE = {
    'mes_referencia': 'Mês de Referência',
    'data_vencimento': 'Data de Vencimento',
    'valor': 'Valor',
    'status': 'Status',
    'data_pagamento': 'Data de Pagamento',
    'valor_pago': 'Valor Pago',
    'observacoes': 'Observações'
}

CAMPOS_PAGAMENTO = {
    'data_pagamento': 'Data do Pagamento',
    'valor': 'Valor',
    'tipo_pagamento': 'Tipo de Pagamento',
    'forma_pagamento': 'Forma de Pagamento',
    'observacoes': 'Observações'
}

CAMPOS_EXTRATO_PIX = {
    'data_pagamento': 'Data do Pagamento',
    'valor': 'Valor',
    'nome_remetente': 'Nome do Remetente',
    'descricao': 'Descrição',
    'status': 'Status'
}

# ==========================================================
# 🤖 FUNÇÕES DE INTELIGÊNCIA ARTIFICIAL
# ==========================================================

def formatar_relatorio_com_ia(dados_brutos: Dict, tipo_relatorio: str, campos_selecionados: List[str]) -> str:
    """
    Usa OpenAI para formatar o relatório de forma inteligente e profissional
    """
    if not openai_client:
        return formatar_relatorio_basico(dados_brutos, tipo_relatorio, campos_selecionados)
    
    try:
        if tipo_relatorio == 'pedagogico':
            exemplo_formato = """
Lista de Alunos

Berçário
1.	Alice Nascimento Rafael
Situação: Matriculado
Responsável Financeiro:
Nome: Mayra Ferreira Nascimento
**OBSERVAÇÃO:**

2.	Ian Duarte Rolim  
Situação: Trancado
Data de Saída: 15/01/2025
Motivo de Saída: Mudança de cidade
Responsável Financeiro:
Nome: Pedro Henrique Rolim de Oliveira
**OBSERVAÇÃO:**
"""
            
            prompt = f"""
Você é um assistente especializado em relatórios pedagógicos para escolas brasileiras.

DADOS RECEBIDOS:
{dados_brutos}

INSTRUÇÕES CRÍTICAS - CAMPOS SELECIONADOS:
RESPEITE RIGOROSAMENTE: Use APENAS os campos selecionados: {', '.join(campos_selecionados)}

CAMPOS DISPONÍVEIS E SEUS NOMES:
ALUNO:
- nome → "Nome do Aluno"
- turno → "Turno" 
- data_nascimento → "Data de Nascimento"
- dia_vencimento → "Dia de Vencimento"
- data_matricula → "Data de Matrícula"
- valor_mensalidade → "Valor da Mensalidade"
- situacao → "Situação"
- mensalidades_geradas → "Mensalidades geradas?"
- data_saida → "Data de Saída" (APENAS se situacao = trancado)
- motivo_saida → "Motivo de Saída" (APENAS se situacao = trancado)

RESPONSÁVEL:
- nome → "Nome do Responsável"
- cpf → "CPF"
- telefone → "Telefone/Contato"
- email → "Email"
- endereco → "Endereço"
- tipo_relacao → "Tipo de Relação"
- responsavel_financeiro → "Responsável Financeiro"

REGRAS DE FORMATAÇÃO:
1. **CAMPOS SELECIONADOS**: Use EXCLUSIVAMENTE os campos da lista: {', '.join(campos_selecionados)}
2. **ORGANIZE POR TURMA**: Lista por turma, alunos numerados em ordem alfabética
3. **CAMPOS VAZIOS**: Para valores NULL, None, vazios ou "Não informado":
   - Use: **_____________** (sublinhado em negrito)
   - NÃO use "AUSENTE"
4. **FORMATAÇÃO DE DADOS**:
   - Datas: DD/MM/YYYY 
   - Valores: R$ X.XXX,XX
   - Situação: Capitalize (Matriculado, Trancado, Problema)
   - Booleanos: "Sim" ou "Não" (mensalidades_geradas, responsavel_financeiro)
5. **CAMPOS ESPECIAIS PARA TRANCADOS**:
   - Se situacao = "trancado" E campos data_saida/motivo_saida foram selecionados → inclua esses campos
   - Se situacao ≠ "trancado" → NÃO inclua data_saida/motivo_saida mesmo se selecionados
6. **RESPONSÁVEIS**:
   - Responsável Financeiro primeiro (se campo selecionado)
   - Outros responsáveis como "Responsável 2, 3..." (se campos selecionados)
   - Inclua APENAS os campos de responsável que foram selecionados
7. **OBSERVAÇÃO**:
   - Inclua "**OBSERVAÇÃO:**" APENAS se houver observações reais para o aluno
   - Se não há observações, NÃO inclua a seção OBSERVAÇÃO
8. **FORMATO DE SAÍDA**:
   - Use ** para negrito
   - Use ** para campos vazios: **_______________**

EXEMPLO DE FORMATO:
{exemplo_formato}

CAMPOS SELECIONADOS A USAR: {', '.join(campos_selecionados)}

IMPORTANTE: 
- NÃO inclua campos não selecionados
- NÃO use "AUSENTE" - use **_______________** para campos vazios
- OBSERVAÇÃO só aparece se houver conteúdo real
- Para alunos trancados, inclua data_saida/motivo_saida APENAS se foram selecionados

Gere o relatório seguindo EXATAMENTE essas regras.
"""
        else:  # financeiro
            # NOVA IMPLEMENTAÇÃO: Relatórios financeiros organizados por status
            status_selecionados = dados_brutos.get('filtros_aplicados', {}).get('status_mensalidades', [])
            tem_campos_responsavel = any(campo in CAMPOS_RESPONSAVEL for campo in campos_selecionados)
            tem_campos_aluno = any(campo in CAMPOS_ALUNO for campo in campos_selecionados)
            tem_campos_mensalidade = any(campo in CAMPOS_MENSALIDADE for campo in campos_selecionados)
            
            # Mapear status para seções - CRÍTICO: 4 seções distintas
            secoes_status = {
                'A vencer': 'MENSALIDADES A VENCER',
                'Pago': 'MENSALIDADES PAGAS',
                'Baixado': 'MENSALIDADES PAGAS', 
                'Pago parcial': 'MENSALIDADES PAGAS',
                'Atrasado': 'MENSALIDADES ATRASADAS',
                'Cancelado': 'MENSALIDADES CANCELADAS'
            }
            
            # Exemplo com múltiplas seções organizadas
            exemplo_financeiro = """
### 1. Alice Nascimento Rafael - Berçário
**Responsáveis:**
💰 **Responsável Financeiro:** Mayra Ferreira Nascimento
   **Telefone:** (83) 99631-0062
   **Email:** ferreiramayra73@gmail.com
👤 **Responsável 2:** Itiel Rafael Figueredo Santos  
   **Telefone:** (83) 99654-6308
   **Email:** **_______________**

**MENSALIDADES PAGAS**
1. **Mês de Referência:** Fevereiro/2025  
   **Data de Vencimento:** 05/02/2025  
   **Data de pagamento:** 05/02/2025
   **Valor mensalidade:** R$ 990,00  
   **Valor pago:** R$ 990,00   
2. **Mês de Referência:** Março/2025  
   **Data de Vencimento:** 05/03/2025  
   **Data de pagamento:** 10/03/2025
   **Valor mensalidade:** R$ 990,00  
   **Valor pago:** R$ 990,00   

**MENSALIDADES A VENCER**
1. **Mês de Referência:** Agosto/2025  
   **Data de Vencimento:** 05/08/2025  
   **Valor:** R$ 990,00  
2. **Mês de Referência:** Setembro/2025  
   **Data de Vencimento:** 05/09/2025  
   **Valor:** R$ 990,00  

---

### 2. João Silva - Infantil I
**Responsáveis:**
💰 **Responsável Financeiro:** Maria Silva
   **Telefone:** (83) 99999-9999

**MENSALIDADES:** NÃO GERADAS

---
"""
            
            prompt = f"""
Você é um assistente especializado em relatórios financeiros organizados por status de mensalidades.

DADOS RECEBIDOS:
{dados_brutos}

INSTRUÇÕES CRÍTICAS - CAMPOS SELECIONADOS:
RESPEITE RIGOROSAMENTE: Use APENAS os campos selecionados: {', '.join(campos_selecionados)}

CAMPOS DISPONÍVEIS E SEUS NOMES:
ALUNO:
- nome → "Nome do Aluno"
- turno → "Turno" 
- data_nascimento → "Data de Nascimento"
- dia_vencimento → "Dia de Vencimento"
- data_matricula → "Data de Matrícula"
- valor_mensalidade → "Valor da Mensalidade"
- situacao → "Situação"
- mensalidades_geradas → "Mensalidades geradas?"

RESPONSÁVEL:
- nome → "Nome do Responsável"
- cpf → "CPF"
- telefone → "Telefone/Contato"
- email → "Email"
- endereco → "Endereço"
- tipo_relacao → "Tipo de Relação"
- responsavel_financeiro → "Responsável Financeiro"

MENSALIDADE:
- mes_referencia → "Mês de Referência"
- data_vencimento → "Data de Vencimento"
- valor → "Valor"
- status → "Status"
- data_pagamento → "Data de Pagamento"
- valor_pago → "Valor Pago"
- observacoes → "Observações"

MAPEAMENTO DE STATUS PARA SEÇÕES (OBRIGATÓRIO - 4 SEÇÕES DISTINTAS):
- "A vencer" → MENSALIDADES A VENCER
- "Pago", "Baixado", "Pago parcial" → MENSALIDADES PAGAS  
- "Atrasado" → MENSALIDADES ATRASADAS
- "Cancelado" → MENSALIDADES CANCELADAS

CRÍTICO: Se status "Pago parcial" existe, DEVE aparecer na seção MENSALIDADES PAGAS

REGRAS DE FORMATAÇÃO POR SEÇÃO:

**PARA MENSALIDADES PAGAS:**
- SEMPRE inclua: Mês de Referência, Data de Vencimento, Data de pagamento, Valor mensalidade, Valor pago
- Formato da data de pagamento: DD/MM/YYYY
- Se data de pagamento for NULL: **_______________**
- Se valor pago for diferente do valor mensalidade, mostre ambos

**PARA OUTRAS SEÇÕES (A vencer, Atrasadas, Canceladas):**
- Inclua: Mês de Referência, Data de Vencimento, Valor
- Formato padrão sem data de pagamento

REGRAS DE ORGANIZAÇÃO:
1. **CADA ALUNO APARECE UMA VEZ** com cabeçalho: "### N. [Nome do Aluno] - [Turma]"
2. **CAMPOS DO ALUNO E RESPONSÁVEIS:** Inclua APENAS os campos selecionados pelo usuário
3. **SEÇÕES DE MENSALIDADES:** 
   - Aparecem APENAS se o status foi selecionado E o aluno tem mensalidades desse status
   - Se nenhuma mensalidade dos tipos selecionados: mostre "**MENSALIDADES:** NÃO GERADAS"
4. **ORDEM DAS SEÇÕES:** PAGAS → A VENCER → ATRASADAS → CANCELADAS
5. **ALUNOS SEM MENSALIDADES:** Mostre "**MENSALIDADES:** NÃO GERADAS"
6. **CAMPOS VAZIOS:** Use **_______________** para NULL/None/vazios (não "AUSENTE")
7. **FORMATAÇÃO:** 
   - Datas: DD/MM/YYYY
   - Valores: R$ X.XXX,XX  
   - Booleanos: "Sim" ou "Não"
   - Use ** para negrito
8. **SEPARAÇÃO:** Use "---" entre alunos
9. **RESPONSÁVEIS:**
   - Use 💰 para responsáveis financeiros
   - Use 👤 para outros responsáveis
   - Inclua TODOS os responsáveis do aluno
   - Para cada responsável, inclua TODOS os campos selecionados

STATUS SELECIONADOS: {status_selecionados}
CAMPOS SELECIONADOS: {', '.join(campos_selecionados)}

ATENÇÃO CRÍTICA: 
- Cada aluno aparece UMA ÚNICA VEZ no relatório
- Use EXCLUSIVAMENTE os dados fornecidos (não invente nomes ou informações)
- Aplique os filtros de status RIGOROSAMENTE
- Se aluno não tem mensalidades dos status selecionados: "MENSALIDADES: NÃO GERADAS"

EXEMPLO DE FORMATO:
{exemplo_financeiro}

IMPORTANTE:
- Ordene mensalidades por data de vencimento (mais antiga primeiro) dentro de cada seção
- Para mensalidades pagas, SEMPRE inclua data de pagamento e valor pago (campos obrigatórios dessa seção)
- Respeite filtros de período aplicados
- Se aluno foi incluído mas não tem mensalidades dos status selecionados, mostre "MENSALIDADES: NÃO GERADAS"
- Use a ordem EXATA dos alunos fornecida

Gere o relatório seguindo EXATAMENTE essas regras.
"""
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Você é um assistente especializado em geração de relatórios educacionais profissionais em português brasileiro."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=16384,
            temperature=0.1
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        print(f"❌ Erro na formatação com IA: {e}")
        return formatar_relatorio_basico(dados_brutos, tipo_relatorio, campos_selecionados)

def formatar_relatorio_basico(dados_brutos: Dict, tipo_relatorio: str, campos_selecionados: List[str]) -> str:
    """
    Formatação básica sem IA como fallback
    """
    texto = f"RELATÓRIO {tipo_relatorio.upper()}\n"
    texto += f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n\n"
    
    if tipo_relatorio == 'pedagogico':
        texto += "Lista de Alunos\n\n"
        
        for turma_nome, dados_turma in dados_brutos.get('dados_por_turma', {}).items():
            texto += f"{turma_nome}\n"
            
            for i, aluno in enumerate(dados_turma.get('alunos', []), 1):
                texto += f"{i}.\t{aluno.get('nome', 'AUSENTE')}\n"
                
                # Adicionar campos selecionados do aluno
                for campo in campos_selecionados:
                    if campo in CAMPOS_ALUNO and campo != 'nome':
                        valor = aluno.get(campo)
                        # Verificar se é realmente NULL/None/vazio
                        if valor is None or valor == "" or valor == "Não informado":
                            valor = 'AUSENTE'
                        elif campo == 'valor_mensalidade' and valor != 'AUSENTE':
                            valor = f"R$ {float(valor):,.2f}"
                        elif campo == 'mensalidades_geradas' and valor != 'AUSENTE':
                            valor = 'Sim' if valor else 'Não'
                        elif 'data' in campo and valor != 'AUSENTE':
                            try:
                                data_obj = datetime.strptime(str(valor), '%Y-%m-%d')
                                valor = data_obj.strftime('%d/%m/%Y')
                            except:
                                pass
                        
                        texto += f"{CAMPOS_ALUNO[campo]}: {valor}\n"
                
                # Adicionar responsáveis se selecionados
                responsaveis = aluno.get('responsaveis', [])
                resp_financeiro = next((r for r in responsaveis if r.get('responsavel_financeiro')), None)
                outros_resp = [r for r in responsaveis if not r.get('responsavel_financeiro')]
                
                # Responsável financeiro primeiro
                if resp_financeiro:
                    texto += "Responsável Financeiro:\n"
                    for campo in campos_selecionados:
                        if campo in CAMPOS_RESPONSAVEL:
                            valor = resp_financeiro.get(campo)
                            # Verificar se é realmente NULL/None/vazio
                            if valor is None or valor == "" or valor == "Não informado":
                                valor = 'AUSENTE'
                            texto += f"{CAMPOS_RESPONSAVEL[campo]}: {valor}\n"
                
                # Outros responsáveis
                for j, resp in enumerate(outros_resp, 2):
                    texto += f"Responsável {j}:\n"
                    for campo in campos_selecionados:
                        if campo in CAMPOS_RESPONSAVEL:
                            valor = resp.get(campo)
                            # Verificar se é realmente NULL/None/vazio
                            if valor is None or valor == "" or valor == "Não informado":
                                valor = 'AUSENTE'
                            texto += f"{CAMPOS_RESPONSAVEL[campo]}: {valor}\n"
                
                # Adicionar campo OBSERVAÇÃO em negrito
                texto += "**OBSERVAÇÃO:**\n\n"
                texto += "\n"
            
            texto += "\n"
    
    elif tipo_relatorio == 'financeiro':
        texto += "Relatório Financeiro Detalhado\n\n"
        
        # Agrupar alunos por turma
        alunos_por_turma = {}
        for aluno in dados_brutos.get('alunos', []):
            turma_nome = aluno.get('turma_nome', 'Sem turma')
            if turma_nome not in alunos_por_turma:
                alunos_por_turma[turma_nome] = []
            alunos_por_turma[turma_nome].append(aluno)
        
        # Mapear mensalidades e pagamentos por aluno
        mensalidades_por_aluno = {}
        for mensalidade in dados_brutos.get('mensalidades', []):
            id_aluno = mensalidade.get('id_aluno')
            if id_aluno not in mensalidades_por_aluno:
                mensalidades_por_aluno[id_aluno] = []
            mensalidades_por_aluno[id_aluno].append(mensalidade)
        
        pagamentos_por_aluno = {}
        for pagamento in dados_brutos.get('pagamentos', []):
            id_aluno = pagamento.get('id_aluno')
            if id_aluno not in pagamentos_por_aluno:
                pagamentos_por_aluno[id_aluno] = []
            pagamentos_por_aluno[id_aluno].append(pagamento)
        
        # Gerar relatório por turma
        for turma_nome, alunos in alunos_por_turma.items():
            texto += f"TURMA: {turma_nome}\n"
            texto += "=" * 50 + "\n\n"
            
            for i, aluno in enumerate(alunos, 1):
                texto += f"{i}. {aluno.get('nome', 'NOME AUSENTE')}\n"
                
                # Dados do aluno
                for campo in campos_selecionados:
                    if campo in CAMPOS_ALUNO and campo != 'nome':
                        valor = aluno.get(campo)
                        if valor is None or valor == "" or valor == "Não informado":
                            valor = 'AUSENTE'
                        elif campo == 'valor_mensalidade' and valor != 'AUSENTE':
                            valor = f"R$ {float(valor):,.2f}"
                        elif campo == 'mensalidades_geradas' and valor != 'AUSENTE':
                            valor = 'Sim' if valor else 'Não'
                        elif 'data' in campo and valor != 'AUSENTE':
                            try:
                                data_obj = datetime.strptime(str(valor), '%Y-%m-%d')
                                valor = data_obj.strftime('%d/%m/%Y')
                            except:
                                pass
                        texto += f"   {CAMPOS_ALUNO[campo]}: {valor}\n"
                
                # Dados dos responsáveis
                responsaveis = aluno.get('responsaveis', [])
                if responsaveis:
                    texto += "   RESPONSÁVEIS:\n"
                    
                    # Separar responsáveis financeiros dos outros
                    resp_financeiros = [r for r in responsaveis if r.get('responsavel_financeiro')]
                    outros_resp = [r for r in responsaveis if not r.get('responsavel_financeiro')]
                    
                    # Primeiro os responsáveis financeiros
                    for j, resp in enumerate(resp_financeiros, 1):
                        emoji = "💰"
                        texto += f"   {emoji} Responsável Financeiro {j}: {resp.get('nome', 'NOME AUSENTE')}\n"
                        
                        for campo in campos_selecionados:
                            if campo in CAMPOS_RESPONSAVEL and campo not in ['nome']:
                                valor = resp.get(campo)
                                if valor is None or valor == "" or valor == "Não informado":
                                    valor = 'AUSENTE'
                                elif campo == 'responsavel_financeiro':
                                    valor = 'SIM' if valor else 'NÃO'
                                texto += f"      {CAMPOS_RESPONSAVEL[campo]}: {valor}\n"
                    
                    # Depois os outros responsáveis
                    for j, resp in enumerate(outros_resp, len(resp_financeiros) + 1):
                        emoji = "👤"
                        texto += f"   {emoji} Responsável {j}: {resp.get('nome', 'NOME AUSENTE')}\n"
                        
                        for campo in campos_selecionados:
                            if campo in CAMPOS_RESPONSAVEL and campo not in ['nome']:
                                valor = resp.get(campo)
                                if valor is None or valor == "" or valor == "Não informado":
                                    valor = 'AUSENTE'
                                elif campo == 'responsavel_financeiro':
                                    valor = 'SIM' if valor else 'NÃO'
                                texto += f"      {CAMPOS_RESPONSAVEL[campo]}: {valor}\n"
                
                # Mensalidades do aluno
                id_aluno = aluno.get('id')
                campos_mensalidade_na_selecao = [campo for campo in campos_selecionados if campo in CAMPOS_MENSALIDADE]
                if campos_mensalidade_na_selecao:
                    mensalidades_aluno = mensalidades_por_aluno.get(id_aluno, [])
                    if mensalidades_aluno:
                        texto += "   MENSALIDADES:\n"
                        for mensalidade in sorted(mensalidades_aluno, key=lambda x: x.get('data_vencimento', '')):
                            for campo in campos_selecionados:
                                if campo in CAMPOS_MENSALIDADE:
                                    valor = mensalidade.get(campo)
                                    if valor is None or valor == "" or valor == "Não informado":
                                        valor = 'AUSENTE'
                                    elif campo == 'valor' and valor != 'AUSENTE':
                                        valor = f"R$ {float(valor):,.2f}"
                                    elif 'data' in campo and valor != 'AUSENTE':
                                        try:
                                            data_obj = datetime.strptime(str(valor), '%Y-%m-%d')
                                            valor = data_obj.strftime('%d/%m/%Y')
                                        except:
                                            pass
                                    texto += f"      {CAMPOS_MENSALIDADE[campo]}: {valor}\n"
                            texto += "      -----\n"
                    else:
                        texto += "   MENSALIDADES: Nenhuma mensalidade encontrada\n"
                
                # Pagamentos do aluno
                campos_pagamento_na_selecao = [campo for campo in campos_selecionados if campo in CAMPOS_PAGAMENTO]
                if campos_pagamento_na_selecao:
                    pagamentos_aluno = pagamentos_por_aluno.get(id_aluno, [])
                    if pagamentos_aluno:
                        texto += "   PAGAMENTOS:\n"
                        for pagamento in sorted(pagamentos_aluno, key=lambda x: x.get('data_pagamento', ''), reverse=True):
                            for campo in campos_selecionados:
                                if campo in CAMPOS_PAGAMENTO:
                                    valor = pagamento.get(campo)
                                    if valor is None or valor == "" or valor == "Não informado":
                                        valor = 'AUSENTE'
                                    elif campo == 'valor' and valor != 'AUSENTE':
                                        valor = f"R$ {float(valor):,.2f}"
                                    elif 'data' in campo and valor != 'AUSENTE':
                                        try:
                                            data_obj = datetime.strptime(str(valor), '%Y-%m-%d')
                                            valor = data_obj.strftime('%d/%m/%Y')
                                        except:
                                            pass
                                    texto += f"      {CAMPOS_PAGAMENTO[campo]}: {valor}\n"
                            texto += "      -----\n"
                    else:
                        texto += "   PAGAMENTOS: Nenhum pagamento encontrado\n"
                
                texto += "\n"
            
            texto += "\n"
        
        # Estatísticas gerais
        texto += "ESTATÍSTICAS GERAIS\n"
        texto += "=" * 30 + "\n"
        total_alunos = len(dados_brutos.get('alunos', []))
        total_mensalidades = len(dados_brutos.get('mensalidades', []))
        total_pagamentos = len(dados_brutos.get('pagamentos', []))
        
        texto += f"Total de Alunos: {total_alunos}\n"
        texto += f"Total de Mensalidades: {total_mensalidades}\n"
        texto += f"Total de Pagamentos: {total_pagamentos}\n"
        
        if dados_brutos.get('filtros_aplicados'):
            texto += "\nFILTROS APLICADOS:\n"
            filtros = dados_brutos['filtros_aplicados']
            for filtro, valor in filtros.items():
                texto += f"  {filtro}: {valor}\n"
    
    return texto

# ==========================================================
# 📊 FUNÇÕES DE COLETA DE DADOS
# ==========================================================

def coletar_dados_pedagogicos(turmas_selecionadas: List[str], campos_selecionados: List[str], 
                             situacoes_filtradas: List[str] = None) -> Dict:
    """
    Coleta dados pedagógicos conforme os filtros selecionados
    
    Args:
        turmas_selecionadas: Lista de nomes das turmas
        campos_selecionados: Lista de campos selecionados pelo usuário
        situacoes_filtradas: Lista de situações para filtrar ['matriculado', 'trancado', 'problema']
    """
    try:
        # Se não especificado, incluir todas as situações
        if not situacoes_filtradas:
            situacoes_filtradas = ["matriculado", "trancado", "problema"]
        
        # Obter IDs das turmas
        mapeamento_resultado = obter_mapeamento_turmas()
        if not mapeamento_resultado.get("success"):
            return {"success": False, "error": "Erro ao obter mapeamento de turmas"}
        
        ids_turmas = []
        for nome_turma in turmas_selecionadas:
            if nome_turma in mapeamento_resultado["mapeamento"]:
                ids_turmas.append(mapeamento_resultado["mapeamento"][nome_turma])
        
        if not ids_turmas:
            return {"success": False, "error": "Nenhuma turma válida selecionada"}
        
        # NOVA IMPLEMENTAÇÃO: Buscar alunos COM FILTRO DE SITUAÇÃO
        dados_organizados = {
            "success": True,
            "dados_por_turma": {},
            "total_alunos": 0,
            "turmas_incluidas": turmas_selecionadas,
            "campos_selecionados": campos_selecionados,
            "situacoes_filtradas": situacoes_filtradas,
            "data_geracao": datetime.now().isoformat()
        }
        
        # Para cada turma, buscar alunos aplicando filtros
        for turma_nome in turmas_selecionadas:
            # Buscar alunos da turma COM filtro de situação
            alunos_response = supabase.table("alunos").select("""
                id, nome, turno, data_nascimento, dia_vencimento, 
                data_matricula, valor_mensalidade, situacao, data_saida, motivo_saida, 
                mensalidades_geradas, turmas!inner(nome_turma)
            """).eq("turmas.nome_turma", turma_nome).in_("situacao", situacoes_filtradas).execute()
            
            if not alunos_response.data:
                # Turma sem alunos ou sem alunos na situação filtrada
                dados_organizados["dados_por_turma"][turma_nome] = {
                    "alunos": [],
                    "total_alunos": 0
                }
                continue
            
            # Ordenar alunos por nome (ordem alfabética)
            alunos_ordenados = sorted(alunos_response.data, key=lambda x: x.get('nome', ''))
            
            alunos_turma = []
            for aluno_data in alunos_ordenados:
                # Buscar responsáveis do aluno
                responsaveis_response = supabase.table("alunos_responsaveis").select("""
                    tipo_relacao, responsavel_financeiro,
                    responsaveis!inner(id, nome, cpf, telefone, email, endereco)
                """).eq("id_aluno", aluno_data["id"]).execute()
                
                # Organizar responsáveis
                responsaveis = []
                for vinculo in responsaveis_response.data:
                    resp_data = vinculo["responsaveis"]
                    resp_data.update({
                        "tipo_relacao": vinculo["tipo_relacao"],
                        "responsavel_financeiro": vinculo["responsavel_financeiro"]
                    })
                    responsaveis.append(resp_data)
                
                # Formatear dados do aluno incluindo novos campos
                aluno_formatado = {
                    "id": aluno_data["id"],
                    "nome": aluno_data["nome"],
                    "turno": aluno_data.get("turno"),
                    "data_nascimento": aluno_data.get("data_nascimento"),
                    "dia_vencimento": aluno_data.get("dia_vencimento"),
                    "data_matricula": aluno_data.get("data_matricula"),
                    "valor_mensalidade": aluno_data.get("valor_mensalidade"),
                    "situacao": aluno_data.get("situacao", "matriculado"),
                    "mensalidades_geradas": aluno_data.get("mensalidades_geradas", False),
                    "data_saida": aluno_data.get("data_saida"),
                    "motivo_saida": aluno_data.get("motivo_saida"),
                    "turma_nome": turma_nome,
                    "responsaveis": responsaveis
                }
                
                alunos_turma.append(aluno_formatado)
            
            dados_organizados["dados_por_turma"][turma_nome] = {
                "alunos": alunos_turma,
                "total_alunos": len(alunos_turma)
            }
            dados_organizados["total_alunos"] += len(alunos_turma)
        
        return dados_organizados
    
    except Exception as e:
        return {"success": False, "error": f"Erro na coleta de dados pedagógicos: {e}"}

def coletar_dados_financeiros(turmas_selecionadas: List[str], campos_selecionados: List[str], 
                             filtros: Dict) -> Dict:
    """
    Coleta dados financeiros conforme os filtros selecionados
    Inclui filtro de situação dos alunos
    """
    try:
        # PRIMEIRO: Atualizar status das mensalidades automaticamente
        try:
            from funcoes_extrato_otimizadas import atualizar_status_mensalidades_automatico
            resultado_atualizacao = atualizar_status_mensalidades_automatico()
            
            # Log da atualização (opcional, pode ser removido se causar ruído)
            if resultado_atualizacao.get("atualizadas", 0) > 0:
                print(f"✅ {resultado_atualizacao['atualizadas']} mensalidades atualizadas para 'Atrasado'")
        except Exception as e:
            print(f"⚠️ Aviso: Erro ao atualizar status automaticamente: {e}")
            # Continua mesmo se a atualização falhar
        
        # ETAPA 1: Buscar alunos das turmas selecionadas COM ORDEM RESPEITADA
        dados_financeiros = {
            "success": True,
            "alunos": [],
            "mensalidades": [],
            "pagamentos": [],
            "extrato_pix": [],
            "filtros_aplicados": filtros,
            "campos_selecionados": campos_selecionados,
            "data_geracao": datetime.now().isoformat()
        }
        
        # ETAPA 2: Para cada turma selecionada, buscar alunos em ordem alfabética
        # Aplicar filtro de situação se especificado
        situacoes_filtradas = filtros.get('situacoes_filtradas', ["matriculado", "trancado", "problema"])
        
        for turma_nome in turmas_selecionadas:
            alunos_response = supabase.table("alunos").select("""
                *, turmas!inner(nome_turma)
            """).eq("turmas.nome_turma", turma_nome).in_("situacao", situacoes_filtradas).execute()
            
            if alunos_response.data:
                # Ordenar alunos por nome (ordem alfabética)
                alunos_ordenados = sorted(alunos_response.data, key=lambda x: x.get('nome', ''))
                
                for aluno_data in alunos_ordenados:
                    aluno_data["turma_nome"] = aluno_data["turmas"]["nome_turma"]
                    
                    # Buscar responsáveis
                    responsaveis_response = supabase.table("alunos_responsaveis").select("""
                        *, responsaveis!inner(*)
                    """).eq("id_aluno", aluno_data["id"]).execute()
                    
                    responsaveis = []
                    for vinculo in responsaveis_response.data:
                        resp_data = vinculo["responsaveis"]
                        resp_data.update({
                            "tipo_relacao": vinculo["tipo_relacao"],
                            "responsavel_financeiro": vinculo["responsavel_financeiro"]
                        })
                        responsaveis.append(resp_data)
                    
                    aluno_data["responsaveis"] = responsaveis
                    dados_financeiros["alunos"].append(aluno_data)
        
        # ETAPA 3: Buscar dados adicionais conforme campos selecionados
        periodo_inicio = filtros.get('periodo_inicio')
        periodo_fim = filtros.get('periodo_fim')
        
        # Coletar IDs de todos os alunos
        ids_alunos = [aluno["id"] for aluno in dados_financeiros["alunos"]]
        
        # Mensalidades - verificar se algum campo de mensalidade foi selecionado
        campos_mensalidade_selecionados = [campo for campo in campos_selecionados if campo in CAMPOS_MENSALIDADE]
        if campos_mensalidade_selecionados and ids_alunos:
            status_mensalidades = filtros.get('status_mensalidades', [])
            
            # CRÍTICO: Filtrar APENAS mensalidades com status especificado
            # Buscar todos os campos incluindo valor_pago
            query = supabase.table("mensalidades").select("*").in_("id_aluno", ids_alunos)
            
            # APLICAR FILTROS OBRIGATÓRIOS
            if status_mensalidades:
                # Converter lista para garantir que apenas os status selecionados sejam incluídos
                query = query.in_("status", status_mensalidades)
            
            # Para relatórios de "Atrasado", adicionar filtro de data de vencimento <= hoje
            if 'Atrasado' in status_mensalidades:
                data_hoje = datetime.now().date().isoformat()
                query = query.lte("data_vencimento", data_hoje)
            
            if periodo_inicio:
                query = query.gte("data_vencimento", periodo_inicio)
            if periodo_fim:
                query = query.lte("data_vencimento", periodo_fim)
            
            mensalidades_response = query.execute()
            
            # NOVA LÓGICA: Manter TODOS os alunos, mas filtrar mensalidades por status
            mensalidades_encontradas = mensalidades_response.data
            
            # Organizar mensalidades por aluno e status para facilitar a IA
            mensalidades_organizadas = {}
            for m in mensalidades_encontradas:
                id_aluno = m.get('id_aluno')
                if id_aluno not in mensalidades_organizadas:
                    mensalidades_organizadas[id_aluno] = {
                        'A vencer': [],
                        'Pago': [],
                        'Baixado': [],
                        'Pago parcial': [],
                        'Atrasado': [],
                        'Cancelado': []
                    }
                
                status = m.get('status', 'A vencer')
                # Agrupar status similares para "MENSALIDADES PAGAS"
                if status in ['Pago', 'Baixado', 'Pago parcial']:
                    # Adicionar à lista correspondente
                    mensalidades_organizadas[id_aluno][status].append(m)
                else:
                    mensalidades_organizadas[id_aluno][status].append(m)
            
            # Adicionar informações de mensalidades organizadas aos alunos
            for aluno in dados_financeiros["alunos"]:
                id_aluno = aluno["id"]
                aluno["mensalidades_por_status"] = mensalidades_organizadas.get(id_aluno, {
                    'A vencer': [], 'Pago': [], 'Baixado': [], 'Pago parcial': [], 'Atrasado': [], 'Cancelado': []
                })
            
            # Manter todas as mensalidades encontradas (não filtrar alunos)
            dados_financeiros["mensalidades"] = mensalidades_encontradas
        else:
            # Se não há campos de mensalidade selecionados, manter alunos mas sem mensalidades
            dados_financeiros["mensalidades"] = []
            for aluno in dados_financeiros["alunos"]:
                aluno["mensalidades_por_status"] = {
                    'A vencer': [], 'Pago': [], 'Baixado': [], 'Pago parcial': [], 'Atrasado': [], 'Cancelado': []
                }
        
        # Pagamentos - verificar se algum campo de pagamento foi selecionado
        campos_pagamento_selecionados = [campo for campo in campos_selecionados if campo in CAMPOS_PAGAMENTO]
        if campos_pagamento_selecionados and ids_alunos:
            query = supabase.table("pagamentos").select("*").in_("id_aluno", ids_alunos)
            
            if periodo_inicio:
                query = query.gte("data_pagamento", periodo_inicio)
            if periodo_fim:
                query = query.lte("data_pagamento", periodo_fim)
            
            pagamentos_response = query.execute()
            dados_financeiros["pagamentos"] = pagamentos_response.data
        
        # Extrato PIX - verificar se algum campo de extrato PIX foi selecionado
        campos_extrato_selecionados = [campo for campo in campos_selecionados if campo in CAMPOS_EXTRATO_PIX]
        if campos_extrato_selecionados and ids_alunos:
            incluir_processados = filtros.get('extrato_pix_processados', False)
            incluir_nao_processados = filtros.get('extrato_pix_nao_processados', False)
            
            # Obter IDs dos responsáveis
            ids_responsaveis = []
            for aluno_data in dados_financeiros["alunos"]:
                for resp in aluno_data.get("responsaveis", []):
                    if resp["id"] not in ids_responsaveis:
                        ids_responsaveis.append(resp["id"])
            
            if ids_responsaveis:
                query = supabase.table("extrato_pix").select("*").in_("id_responsavel", ids_responsaveis)
                
                if periodo_inicio:
                    query = query.gte("data_pagamento", periodo_inicio)
                if periodo_fim:
                    query = query.lte("data_pagamento", periodo_fim)
                
                status_filtros = []
                if incluir_processados:
                    status_filtros.append("registrado")
                if incluir_nao_processados:
                    status_filtros.append("novo")
                
                if status_filtros:
                    query = query.in_("status", status_filtros)
                
                extrato_response = query.execute()
                dados_financeiros["extrato_pix"] = extrato_response.data
        
        return dados_financeiros
    
    except Exception as e:
        return {"success": False, "error": f"Erro na coleta de dados financeiros: {e}"}

# ==========================================================
# 📄 FUNÇÕES DE GERAÇÃO DE DOCUMENTOS
# ==========================================================

def criar_documento_docx(titulo: str, conteudo: str) -> Optional[Document]:
    """
    Cria um documento .docx formatado profissionalmente
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Configurar margens (A4, orientação vertical)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título principal
        titulo_para = doc.add_heading(titulo, level=1)
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data de geração
        data_para = doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Separador
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar conteúdo processando formatação em negrito
        linhas = conteudo.split('\n')
        for linha in linhas:
            if linha.strip():
                if linha.strip().endswith(':') and len(linha) < 50 and not '**' in linha:
                    # Títulos de seção (sem ** no texto)
                    para = doc.add_paragraph()
                    run = para.add_run(linha)
                    run.bold = True
                elif linha.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and not '**' in linha:
                    # Numeração de alunos (sem ** no texto)
                    para = doc.add_paragraph()
                    run = para.add_run(linha)
                    run.bold = True
                    run.font.size = Pt(12)
                else:
                    # Processar linha com possível formatação em negrito
                    para = doc.add_paragraph()
                    processar_linha_com_negrito(para, linha)
            else:
                doc.add_paragraph()
        
        return doc
    
    except Exception as e:
        print(f"❌ Erro ao criar documento: {e}")
        return None

def processar_linha_com_negrito(paragrafo, texto: str):
    """
    Processa uma linha de texto, aplicando formatação em negrito para texto entre **
    Inclui suporte especial para campos vazios formatados como **_______________**
    """
    import re
    
    # Primeiro processar campos vazios especiais **_______________**
    padrao_vazio = r'\*\*(_+)\*\*'
    
    # Dividir por campos vazios primeiro
    partes_vazio = re.split(padrao_vazio, texto)
    
    for i, parte_vazio in enumerate(partes_vazio):
        if parte_vazio and parte_vazio.startswith('___'):
            # É um campo vazio - criar run sublinhado em negrito e vermelho
            run = paragrafo.add_run(parte_vazio)
            run.bold = True
            run.underline = True
            # Definir cor vermelha (se suportado)
            try:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho
            except:
                pass  # Se não conseguir aplicar cor, continua sem
        elif parte_vazio:
            # Texto normal - processar padrão de negrito normal
            padrao_negrito = r'\*\*(.*?)\*\*'
            partes_negrito = re.split(padrao_negrito, parte_vazio)
            
            for j, parte in enumerate(partes_negrito):
                if parte:  # Ignorar strings vazias
                    run = paragrafo.add_run(parte)
                    
                    # Se é uma parte ímpar, estava entre ** então deve ficar em negrito
                    if j % 2 == 1:
                        run.bold = True

def salvar_documento_temporario(doc: Document, nome_arquivo: str) -> Optional[str]:
    """
    Salva documento temporariamente e retorna o caminho
    """
    try:
        pasta_temp = "temp_reports"
        if not os.path.exists(pasta_temp):
            os.makedirs(pasta_temp)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_completo = f"{nome_arquivo}_{timestamp}.docx"
        caminho_arquivo = os.path.join(pasta_temp, nome_completo)
        
        doc.save(caminho_arquivo)
        return caminho_arquivo
    
    except Exception as e:
        print(f"❌ Erro ao salvar documento: {e}")
        return None

# ==========================================================
# 🎯 FUNÇÕES PRINCIPAIS DE GERAÇÃO
# ==========================================================

def gerar_relatorio_pedagogico(turmas_selecionadas: List[str], campos_selecionados: List[str], 
                              situacoes_filtradas: List[str] = None) -> Dict:
    """
    Gera relatório pedagógico completo
    
    Args:
        turmas_selecionadas: Lista de turmas para incluir
        campos_selecionados: Lista de campos para exibir
        situacoes_filtradas: Lista de situações para filtrar ['matriculado', 'trancado', 'problema']
    """
    try:
        if not DOCX_AVAILABLE:
            return {
                "success": False, 
                "error": "python-docx não disponível. Execute: pip install python-docx"
            }
        
        # Coletar dados COM filtro de situação
        dados = coletar_dados_pedagogicos(turmas_selecionadas, campos_selecionados, situacoes_filtradas)
        if not dados.get("success"):
            return dados
        
        # Formatar com IA
        conteudo_formatado = formatar_relatorio_com_ia(dados, "pedagogico", campos_selecionados)
        
        # Criar documento
        titulo = f"Relatório Pedagógico - {', '.join(turmas_selecionadas)}"
        doc = criar_documento_docx(titulo, conteudo_formatado)
        
        if not doc:
            return {"success": False, "error": "Erro ao criar documento Word"}
        
        # Salvar temporariamente
        nome_arquivo = f"relatorio_pedagogico_{'_'.join(turmas_selecionadas)}"
        nome_arquivo = nome_arquivo.replace(' ', '_').replace('/', '_')
        
        caminho_arquivo = salvar_documento_temporario(doc, nome_arquivo)
        
        if not caminho_arquivo:
            return {"success": False, "error": "Erro ao salvar documento"}
        
        return {
            "success": True,
            "arquivo": caminho_arquivo,
            "nome_arquivo": os.path.basename(caminho_arquivo),
            "titulo": titulo,
            "total_alunos": dados["total_alunos"],
            "turmas_incluidas": dados["turmas_incluidas"],
            "campos_selecionados": dados["campos_selecionados"],
            "situacoes_filtradas": dados["situacoes_filtradas"],
            "data_geracao": dados["data_geracao"]
        }
    
    except Exception as e:
        return {"success": False, "error": f"Erro na geração do relatório pedagógico: {e}"}

def gerar_relatorio_financeiro(turmas_selecionadas: List[str], campos_selecionados: List[str], 
                              filtros: Dict) -> Dict:
    """
    Gera relatório financeiro completo
    """
    try:
        if not DOCX_AVAILABLE:
            return {
                "success": False, 
                "error": "python-docx não disponível. Execute: pip install python-docx"
            }
        
        # Coletar dados
        dados = coletar_dados_financeiros(turmas_selecionadas, campos_selecionados, filtros)
        if not dados.get("success"):
            return dados
        
        # Formatar com IA
        conteudo_formatado = formatar_relatorio_com_ia(dados, "financeiro", campos_selecionados)
        
        # Criar documento
        titulo = f"Relatório Financeiro - {', '.join(turmas_selecionadas)}"
        doc = criar_documento_docx(titulo, conteudo_formatado)
        
        if not doc:
            return {"success": False, "error": "Erro ao criar documento Word"}
        
        # Salvar temporariamente
        nome_arquivo = f"relatorio_financeiro_{'_'.join(turmas_selecionadas)}"
        nome_arquivo = nome_arquivo.replace(' ', '_').replace('/', '_')
        
        caminho_arquivo = salvar_documento_temporario(doc, nome_arquivo)
        
        if not caminho_arquivo:
            return {"success": False, "error": "Erro ao salvar documento"}
        
        return {
            "success": True,
            "arquivo": caminho_arquivo,
            "arquivo_temporario": caminho_arquivo,  # Alias para o teste
            "nome_arquivo": os.path.basename(caminho_arquivo),
            "titulo": titulo,
            "conteudo_formatado": conteudo_formatado,  # Adicionar conteúdo para análise
            "total_alunos": len(dados.get("alunos", [])),
            "turmas_incluidas": turmas_selecionadas,
            "campos_selecionados": campos_selecionados,
            "filtros_aplicados": filtros,
            "data_geracao": dados["data_geracao"]
        }
    
    except Exception as e:
        return {"success": False, "error": f"Erro na geração do relatório financeiro: {e}"}

# ==========================================================
# 🧹 FUNÇÕES AUXILIARES
# ==========================================================

def limpar_arquivos_temporarios(max_idade_horas: int = 24):
    """Remove arquivos temporários antigos"""
    try:
        pasta_temp = "temp_reports"
        if not os.path.exists(pasta_temp):
            return
        
        agora = datetime.now()
        arquivos_removidos = 0
        
        for arquivo in os.listdir(pasta_temp):
            caminho_arquivo = os.path.join(pasta_temp, arquivo)
            tempo_modificacao = datetime.fromtimestamp(os.path.getmtime(caminho_arquivo))
            idade_horas = (agora - tempo_modificacao).total_seconds() / 3600
            
            if idade_horas > max_idade_horas:
                os.remove(caminho_arquivo)
                arquivos_removidos += 1
        
        if arquivos_removidos > 0:
            print(f"🧹 {arquivos_removidos} arquivos temporários removidos")
    
    except Exception as e:
        print(f"❌ Erro ao limpar arquivos temporários: {e}")

def obter_campos_disponiveis() -> Dict:
    """Retorna todos os campos disponíveis para relatórios"""
    return {
        "aluno": CAMPOS_ALUNO,
        "responsavel": CAMPOS_RESPONSAVEL,
        "mensalidade": CAMPOS_MENSALIDADE,
        "pagamento": CAMPOS_PAGAMENTO,
        "extrato_pix": CAMPOS_EXTRATO_PIX
    }

# ==========================================================
# 🎯 FUNÇÃO PRINCIPAL PARA INTERFACE
# ==========================================================

def gerar_relatorio_interface(tipo_relatorio: str, configuracao: Dict) -> Dict:
    """
    Função principal para ser chamada pela interface Streamlit
    """
    try:
        # Validar tipo de relatório
        if tipo_relatorio not in ['pedagogico', 'financeiro']:
            return {"success": False, "error": "Tipo de relatório inválido"}
        
        # Extrair configurações
        turmas_selecionadas = configuracao.get('turmas_selecionadas', [])
        campos_selecionados = configuracao.get('campos_selecionados', [])
        
        # Validar turmas
        if not turmas_selecionadas:
            return {"success": False, "error": "Nenhuma turma selecionada"}
        
        # Validar campos
        if not campos_selecionados:
            return {"success": False, "error": "Nenhum campo selecionado"}
        
        # Limpar arquivos antigos
        limpar_arquivos_temporarios()
        
        # Gerar relatório conforme o tipo
        if tipo_relatorio == 'pedagogico':
            situacoes_filtradas = configuracao.get('situacoes_filtradas', [])
            return gerar_relatorio_pedagogico(turmas_selecionadas, campos_selecionados, situacoes_filtradas)
        else:
            filtros = configuracao.get('filtros', {})
            return gerar_relatorio_financeiro(turmas_selecionadas, campos_selecionados, filtros)
    
    except Exception as e:
        return {"success": False, "error": f"Erro na geração do relatório: {e}"} 